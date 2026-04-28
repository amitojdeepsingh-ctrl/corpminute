"""
Document generation engine for CorpMinute.ca
Generates all 10 minute book documents + special resolutions using python-docx
"""
import io
import zipfile
from datetime import datetime, date
from pathlib import Path
from typing import Optional

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schema import Corporation, Director, Officer, Shareholder
from documents import DISCLAIMER, JURISDICTION_CITATIONS, PROVINCE_NAMES


def _styled_doc() -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    return doc


def _add_header(doc: Document, corp_name: str, title: str) -> None:
    header = doc.sections[0].header
    p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    p.clear()
    run = p.add_run(f"{corp_name}  |  {title}")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _add_footer_disclaimer(doc: Document) -> None:
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    run = p.add_run(DISCLAIMER)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_title_block(doc: Document, corp_name: str, doc_title: str, province: str) -> None:
    doc.add_paragraph()
    h = doc.add_heading(corp_name.upper(), level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)

    sub = doc.add_paragraph(doc_title)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(14)
    sub.runs[0].bold = True

    prov = doc.add_paragraph(f"{PROVINCE_NAMES.get(province, province)} Corporation")
    prov.alignment = WD_ALIGN_PARAGRAPH.CENTER
    prov.runs[0].font.size = Pt(11)

    cite = doc.add_paragraph(JURISDICTION_CITATIONS.get(province, ""))
    cite.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cite.runs[0].font.size = Pt(9)
    cite.runs[0].italic = True
    cite.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()


def _today_str() -> str:
    return datetime.utcnow().strftime("%B %d, %Y")


def _fiscal_year_display(fiscal_year_end: str) -> str:
    if not fiscal_year_end:
        return "December 31"
    try:
        d = datetime.strptime(fiscal_year_end, "%Y-%m-%d")
        return d.strftime("%B %d")
    except Exception:
        return fiscal_year_end


# ─────────────────────────────────────────────
# DOCUMENT 1: Annual Directors' Resolution
# ─────────────────────────────────────────────
def generate_directors_resolution(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Annual Directors' Resolution")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "ANNUAL RESOLUTION OF THE DIRECTORS", corp.province)

    doc.add_paragraph(
        f"The undersigned, being all of the directors of {corp.corp_name} "
        f"(the \"Corporation\"), a corporation incorporated under the laws of "
        f"{PROVINCE_NAMES.get(corp.province, corp.province)}, hereby consent to the "
        f"passing of the following resolution in lieu of a meeting, pursuant to "
        f"{JURISDICTION_CITATIONS.get(corp.province, 'applicable legislation')}:"
    )
    doc.add_paragraph()

    # Resolution body
    doc.add_heading("RESOLVED THAT:", level=3)
    resolutions = [
        f"1. The financial statements of the Corporation for the fiscal year ended "
        f"{_fiscal_year_display(corp.fiscal_year_end)} are hereby approved and adopted.",
        "2. The audit of the financial statements of the Corporation is hereby waived for "
        "the current fiscal year, as permitted by the applicable corporations legislation.",
        "3. The current officer(s) of the Corporation are hereby confirmed in their respective "
        "roles until the next annual meeting of the directors or until their successors are "
        "duly appointed.",
        "4. The directors are hereby authorized to transact all business of the Corporation "
        "in the ordinary course.",
    ]
    for r in resolutions:
        doc.add_paragraph(r)

    doc.add_paragraph()
    doc.add_paragraph(f"DATED this {_today_str()}.")
    doc.add_paragraph()

    for director in corp.directors:
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph(f"{director.name}, Director")

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 2: Annual Shareholders' Resolution
# ─────────────────────────────────────────────
def generate_shareholders_resolution(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Annual Shareholders' Resolution")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "ANNUAL RESOLUTION OF THE SHAREHOLDERS", corp.province)

    doc.add_paragraph(
        f"The undersigned, being all of the shareholders of {corp.corp_name} "
        f"entitled to vote at a meeting of shareholders, hereby consent to the passing "
        f"of the following resolution in lieu of a meeting, pursuant to "
        f"{JURISDICTION_CITATIONS.get(corp.province, 'applicable legislation')}:"
    )
    doc.add_paragraph()

    doc.add_heading("RESOLVED THAT:", level=3)

    director_names = " and ".join(d.name for d in corp.directors) if corp.directors else "[Director Names]"
    resolutions = [
        f"1. {director_names} {'is' if len(corp.directors) == 1 else 'are'} hereby re-elected "
        f"as {'director' if len(corp.directors) == 1 else 'directors'} of the Corporation to "
        f"hold office until the next annual meeting of shareholders or until "
        f"{'a successor is' if len(corp.directors) == 1 else 'successors are'} duly elected or appointed.",
        "2. The appointment of an auditor is hereby waived for the current fiscal year, "
        "as all shareholders have consented to such waiver in accordance with the applicable "
        "corporations legislation.",
        "3. Any one director or officer of the Corporation is hereby authorized to execute "
        "and deliver all documents and to take all steps necessary or desirable to give "
        "effect to the foregoing resolutions.",
    ]
    for r in resolutions:
        doc.add_paragraph(r)

    doc.add_paragraph()
    doc.add_paragraph(f"DATED this {_today_str()}.")
    doc.add_paragraph()

    for shareholder in corp.shareholders:
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph(f"{shareholder.name}, Shareholder ({shareholder.quantity} {shareholder.share_class} shares)")

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 3: Register of Directors
# ─────────────────────────────────────────────
def generate_register_of_directors(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Register of Directors")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "REGISTER OF DIRECTORS", corp.province)

    doc.add_paragraph(
        f"This Register of Directors is maintained pursuant to "
        f"{JURISDICTION_CITATIONS.get(corp.province, 'applicable legislation')}. "
        f"Last updated: {_today_str()}."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Full Name", "Residential Address", "Date Appointed", "Date Ceased"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True

    for d in corp.directors:
        row = table.add_row()
        row.cells[0].text = d.name
        row.cells[1].text = d.address
        row.cells[2].text = d.appointed
        row.cells[3].text = ""

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 4: Register of Officers
# ─────────────────────────────────────────────
def generate_register_of_officers(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Register of Officers")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "REGISTER OF OFFICERS", corp.province)

    doc.add_paragraph(
        f"This Register of Officers is maintained pursuant to "
        f"{JURISDICTION_CITATIONS.get(corp.province, 'applicable legislation')}. "
        f"Last updated: {_today_str()}."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Full Name", "Office Held", "Date Appointed", "Date Ceased"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for o in corp.officers:
        row = table.add_row()
        row.cells[0].text = o.name
        row.cells[1].text = o.role
        row.cells[2].text = o.appointed
        row.cells[3].text = ""

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 5: Register of Shareholders
# ─────────────────────────────────────────────
def generate_register_of_shareholders(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Register of Shareholders")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "REGISTER OF SHAREHOLDERS", corp.province)

    doc.add_paragraph(
        f"This Register of Shareholders is maintained pursuant to "
        f"{JURISDICTION_CITATIONS.get(corp.province, 'applicable legislation')}. "
        f"Last updated: {_today_str()}."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    headers = ["Full Name / Entity", "Share Class", "Shares Held", "Address"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for s in corp.shareholders:
        row = table.add_row()
        row.cells[0].text = s.name
        row.cells[1].text = s.share_class
        row.cells[2].text = str(s.quantity)
        row.cells[3].text = ""

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 6: Share Ledger
# ─────────────────────────────────────────────
def generate_share_ledger(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Share Ledger")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "SHARE LEDGER", corp.province)

    doc.add_paragraph(
        f"This Share Ledger records all share issuances and transfers of {corp.corp_name}. "
        f"Last updated: {_today_str()}."
    )
    doc.add_paragraph()

    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ["Certificate #", "Shareholder", "Class", "Shares", "Date Issued"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    for idx, s in enumerate(corp.shareholders, start=1):
        row = table.add_row()
        row.cells[0].text = f"C-{idx:04d}"
        row.cells[1].text = s.name
        row.cells[2].text = s.share_class
        row.cells[3].text = str(s.quantity)
        row.cells[4].text = corp.incorporation_date or _today_str()

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 7: Consent to Act as Director
# ─────────────────────────────────────────────
def generate_consent_to_act(corp: Corporation, director: Director) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Consent to Act as Director")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "CONSENT TO ACT AS DIRECTOR", corp.province)

    doc.add_paragraph(
        f"I, {director.name}, of {director.address}, hereby consent to act as a director "
        f"of {corp.corp_name}, a corporation incorporated under the laws of "
        f"{PROVINCE_NAMES.get(corp.province, corp.province)}, pursuant to "
        f"{JURISDICTION_CITATIONS.get(corp.province, 'applicable legislation')}."
    )
    doc.add_paragraph()
    doc.add_paragraph(f"DATED this {_today_str()}.")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph("_" * 40)
    doc.add_paragraph(f"{director.name}")
    doc.add_paragraph("Director")

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 8: Organizational Resolution
# ─────────────────────────────────────────────
def generate_organizational_resolution(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Organizational Resolution")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "ORGANIZATIONAL RESOLUTION OF THE DIRECTORS", corp.province)

    doc.add_paragraph(
        f"The directors of {corp.corp_name} hereby pass the following organizational "
        f"resolution in lieu of an organizational meeting, pursuant to "
        f"{JURISDICTION_CITATIONS.get(corp.province, 'applicable legislation')}:"
    )
    doc.add_paragraph()

    doc.add_heading("RESOLVED THAT:", level=3)

    officer_name = corp.officers[0].name if corp.officers else (corp.directors[0].name if corp.directors else "[Officer Name]")

    resolutions = [
        f"1. {officer_name} is hereby appointed President and Secretary of the Corporation.",
        f"2. The fiscal year of the Corporation shall end on {_fiscal_year_display(corp.fiscal_year_end)} "
        f"of each year.",
        "3. The Corporation is hereby authorized to open one or more bank accounts and to conduct "
        "banking business as the directors may determine.",
        "4. The by-laws of the Corporation, if any, are hereby confirmed and adopted.",
        "5. Any director or officer is hereby authorized to execute any documents necessary to "
        "complete the organization of the Corporation.",
    ]
    for r in resolutions:
        doc.add_paragraph(r)

    doc.add_paragraph()
    doc.add_paragraph(f"DATED this {_today_str()}.")
    doc.add_paragraph()

    for director in corp.directors:
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph(f"{director.name}, Director")

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 9: Banking Resolution Template
# ─────────────────────────────────────────────
def generate_banking_resolution(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_header(doc, corp.corp_name, "Banking Resolution")
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, "BANKING RESOLUTION OF THE DIRECTORS", corp.province)

    doc.add_paragraph(
        f"The directors of {corp.corp_name} hereby resolve as follows:"
    )
    doc.add_paragraph()

    doc.add_heading("RESOLVED THAT:", level=3)

    officer_name = corp.officers[0].name if corp.officers else (corp.directors[0].name if corp.directors else "[Officer Name]")

    resolutions = [
        "1. The Corporation is hereby authorized to open and maintain one or more bank accounts "
        "at such financial institution(s) as the directors may determine.",
        f"2. {officer_name}, or such other persons as the directors may designate from time to time, "
        "are hereby authorized to operate such bank accounts, including signing cheques, making "
        "electronic transfers, and executing any documentation required by the financial institution.",
        "3. The financial institution is hereby authorized to honour all cheques, drafts, and "
        "orders signed by the authorized signatories as set out above.",
        "4. A copy of this resolution, certified by any director or officer, shall constitute "
        "sufficient authority for any financial institution to act upon.",
    ]
    for r in resolutions:
        doc.add_paragraph(r)

    doc.add_paragraph()
    doc.add_paragraph(f"DATED this {_today_str()}.")
    doc.add_paragraph()

    for director in corp.directors:
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph(f"{director.name}, Director")

    return doc


# ─────────────────────────────────────────────
# DOCUMENT 10: Cover Page + Table of Contents
# ─────────────────────────────────────────────
def generate_cover_page(corp: Corporation) -> Document:
    doc = _styled_doc()
    _add_footer_disclaimer(doc)

    # Spacer
    for _ in range(6):
        doc.add_paragraph()

    h = doc.add_heading("MINUTE BOOK", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x6B)
        run.font.size = Pt(28)

    corp_p = doc.add_paragraph(corp.corp_name.upper())
    corp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    corp_p.runs[0].font.size = Pt(18)
    corp_p.runs[0].bold = True

    prov_p = doc.add_paragraph(PROVINCE_NAMES.get(corp.province, corp.province))
    prov_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    corp_num_p = doc.add_paragraph(f"Corporation No. {corp.corp_number or '[To be inserted]'}")
    corp_num_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    inc_p = doc.add_paragraph(f"Incorporated: {corp.incorporation_date or '[Date]'}")
    inc_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for _ in range(4):
        doc.add_paragraph()

    gen_p = doc.add_paragraph(f"Generated by CorpMinute.ca  |  {_today_str()}")
    gen_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_p.runs[0].font.size = Pt(9)
    gen_p.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    doc.add_page_break()

    # Table of Contents
    toc_heading = doc.add_heading("TABLE OF CONTENTS", level=2)
    toc_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

    toc_items = [
        ("1.", "Annual Directors' Resolution"),
        ("2.", "Annual Shareholders' Resolution"),
        ("3.", "Register of Directors"),
        ("4.", "Register of Officers"),
        ("5.", "Register of Shareholders"),
        ("6.", "Share Ledger"),
        ("7.", "Consent to Act as Director"),
        ("8.", "Organizational Resolution"),
        ("9.", "Banking Resolution"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        p.add_run(f"  {num}  ").bold = True
        p.add_run(title)

    return doc


# ─────────────────────────────────────────────
# SPECIAL RESOLUTIONS
# ─────────────────────────────────────────────
def generate_special_resolution(
    corp: Corporation,
    resolution_type: str,
    details: str,
    resolution_date: Optional[str] = None,
) -> Document:
    doc = _styled_doc()
    res_date = resolution_date or _today_str()

    SPECIAL_TYPES = {
        "bank_account": "BANKING AUTHORIZATION RESOLUTION",
        "share_issuance": "RESOLUTION TO ISSUE SHARES",
        "new_director": "RESOLUTION APPOINTING NEW DIRECTOR",
        "loan_authorization": "LOAN / LINE OF CREDIT AUTHORIZATION RESOLUTION",
        "fiscal_year_change": "RESOLUTION TO CHANGE FISCAL YEAR-END",
        "amend_articles": "RESOLUTION TO AMEND ARTICLES",
        "officer_change": "OFFICER APPOINTMENT / REMOVAL RESOLUTION",
    }

    title = SPECIAL_TYPES.get(resolution_type, "SPECIAL RESOLUTION")
    _add_header(doc, corp.corp_name, title)
    _add_footer_disclaimer(doc)
    _add_title_block(doc, corp.corp_name, title, corp.province)

    doc.add_paragraph(
        f"The directors of {corp.corp_name} hereby resolve as follows:"
    )
    doc.add_paragraph()
    doc.add_heading("RESOLVED THAT:", level=3)
    doc.add_paragraph(details)
    doc.add_paragraph()
    doc.add_paragraph(f"DATED this {res_date}.")
    doc.add_paragraph()

    for director in corp.directors:
        doc.add_paragraph()
        doc.add_paragraph("_" * 40)
        doc.add_paragraph(f"{director.name}, Director")

    return doc


# ─────────────────────────────────────────────
# FULL MINUTE BOOK — generate all 10 docs
# ─────────────────────────────────────────────
def generate_full_minute_book(corp: Corporation) -> dict[str, Document]:
    """Returns a dict of filename -> Document for all 10 documents."""
    docs = {}
    docs["00_cover.docx"] = generate_cover_page(corp)
    docs["01_directors_resolution.docx"] = generate_directors_resolution(corp)
    docs["02_shareholders_resolution.docx"] = generate_shareholders_resolution(corp)
    docs["03_register_of_directors.docx"] = generate_register_of_directors(corp)
    docs["04_register_of_officers.docx"] = generate_register_of_officers(corp)
    docs["05_register_of_shareholders.docx"] = generate_register_of_shareholders(corp)
    docs["06_share_ledger.docx"] = generate_share_ledger(corp)
    docs["07_banking_resolution.docx"] = generate_banking_resolution(corp)
    docs["08_organizational_resolution.docx"] = generate_organizational_resolution(corp)

    for idx, director in enumerate(corp.directors, start=1):
        docs[f"09_consent_director_{idx}.docx"] = generate_consent_to_act(corp, director)

    return docs


def docs_to_zip(docs: dict[str, Document]) -> bytes:
    """Bundle all documents into a ZIP file in memory."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for filename, doc in docs.items():
            doc_buf = io.BytesIO()
            doc.save(doc_buf)
            zf.writestr(filename, doc_buf.getvalue())
    return buf.getvalue()


def save_docs_to_dir(docs: dict[str, Document], output_dir: Path) -> None:
    output_dir.mkdir(parents=True, exist_ok=True)
    for filename, doc in docs.items():
        doc.save(output_dir / filename)
